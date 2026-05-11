from typing import List, Dict, Any, Optional
from langchain_core.tools import tool
from langchain_office_assistant.plugins.base import BasePlugin
from langchain_office_assistant.utils.logger import get_logger
import docx
import openpyxl
from PyPDF2 import PdfReader
import os

logger = get_logger(__name__)

class DocumentPlugin(BasePlugin):
    name = "document"
    description = "文档处理插件 - Word/Excel/PDF读写"
    
    def __init__(self):
        super().__init__()
        self.documents = [
            {
                "id": "doc_001",
                "name": "项目报告.docx",
                "type": "docx",
                "path": "documents/project_report.docx",
                "size": 10240,
                "modified": "2025-01-14"
            },
            {
                "id": "doc_002",
                "name": "会议记录.pdf",
                "type": "pdf",
                "path": "documents/meeting_notes.pdf",
                "size": 5120,
                "modified": "2025-01-13"
            },
            {
                "id": "doc_003",
                "name": "数据报表.xlsx",
                "type": "xlsx",
                "path": "documents/data_report.xlsx",
                "size": 8192,
                "modified": "2025-01-12"
            }
        ]
    
    def initialize(self, config: Dict) -> None:
        self.config = config
        logger.info(f"DocumentPlugin initialized")
    
    def get_tools(self) -> List:
        return [search_documents, summarize_document, read_document, write_document]
    
    async def execute(self, tool_name: str, **kwargs) -> Any:
        tools_map = {
            "search_documents": search_documents,
            "summarize_document": summarize_document,
            "read_document": read_document,
            "write_document": write_document,
        }
        
        if tool_name not in tools_map:
            return f"❌ Tool not found: {tool_name}"
        
        tool_func = tools_map[tool_name]
        return tool_func(**kwargs)

@tool
def search_documents(keyword: str) -> str:
    mock_results = [
        {
            "id": "doc_001",
            "name": "项目报告.docx",
            "type": "docx",
            "preview": f"文档内容包含 '{keyword}' 相关章节..."
        },
        {
            "id": "doc_002", 
            "name": "会议记录.pdf",
            "type": "pdf",
            "preview": f"会议讨论了与 '{keyword}' 相关的内容..."
        }
    ]
    
    output = f"🔍 Search results for '{keyword}':\n\n"
    for doc in mock_results:
        output += (
            f"📄 {doc['name']} ({doc['type']})\n"
            f"   Preview: {doc['preview']}\n"
            f"   ID: {doc['id']}\n\n"
        )
    
    return output

@tool
def summarize_document(file_path: str) -> str:
    try:
        ext = file_path.split('.')[-1].lower()
        
        if ext == 'docx':
            doc = docx.Document(file_path)
            content = '\n'.join([p.text for p in doc.paragraphs])
        elif ext == 'pdf':
            reader = PdfReader(file_path)
            content = '\n'.join([page.extract_text() for page in reader.pages])
        elif ext == 'xlsx':
            wb = openpyxl.load_workbook(file_path)
            content = ""
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                content += f"Sheet: {sheet}\n"
                for row in ws.iter_rows(values_only=True):
                    content += f"{row}\n"
        else:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        
        lines = content.split('\n')[:5]
        preview = '\n'.join(lines)
        
        return (
            f"📄 Document Summary: {file_path}\n\n"
            f"Key Points:\n"
            f"1. Document type: {ext.upper()}\n"
            f"2. Content preview:\n{preview}\n"
            f"3. Full content length: {len(content)} characters\n"
            f"4. Contains {len(content.split())} words"
        )
    except Exception as e:
        return f"❌ Failed to summarize document: {str(e)}"

@tool
def read_document(file_path: str) -> str:
    try:
        ext = file_path.split('.')[-1].lower()
        
        if ext == 'docx':
            doc = docx.Document(file_path)
            content = '\n'.join([p.text for p in doc.paragraphs])
        elif ext == 'pdf':
            reader = PdfReader(file_path)
            content = '\n'.join([page.extract_text() for page in reader.pages])
        elif ext == 'xlsx':
            wb = openpyxl.load_workbook(file_path)
            content = ""
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                content += f"=== Sheet: {sheet} ===\n"
                for row in ws.iter_rows(values_only=True):
                    content += f"{row}\n"
        else:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        
        return f"📄 Content of {file_path}:\n\n{content}"
    except Exception as e:
        return f"❌ Failed to read document: {str(e)}"

@tool
def write_document(file_path: str, content: str, format: str = "txt") -> str:
    try:
        ext = format.lower()
        
        if ext == 'docx':
            doc = docx.Document()
            doc.add_paragraph(content)
            doc.save(file_path)
        elif ext == 'xlsx':
            wb = openpyxl.Workbook()
            ws = wb.active
            for i, line in enumerate(content.split('\n'), 1):
                ws.cell(row=i, column=1, value=line)
            wb.save(file_path)
        else:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
        
        return f"✅ Document saved successfully!\n\n📄 File: {file_path}\n📊 Format: {format.upper()}"
    except Exception as e:
        return f"❌ Failed to write document: {str(e)}"